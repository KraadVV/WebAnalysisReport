# src/report_writer.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


class DocxGenerator:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        '''문서의 기본 폰트나 스타일 설정 (선택사항)'''
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'  # 한글 폰트 설정 (시스템에 있어야 함)
        font.size = Pt(10)

    def create_report(self, url, screenshot_path, page_data, api_logs, interactive_elements=None, whois_data=None, osint_data=None, unique_ips=None, resource_summary=None, sns_data=None):
        '''
        모든 데이터를 받아서 순서대로 문서에 작성.
        '''

        # 1. 제목 및 기본 정보
        self.doc.add_heading(f'웹사이트 분석 보고서', 0)

        p = self.doc.add_paragraph()
        p.add_run(f"분석 대상 URL: ").bold = True
        p.add_run(url + "\n")
        p.add_run(f"사이트 제목: ").bold = True
        p.add_run(page_data.get('title', 'N/A'))

        # 2. 초기 화면 스크린샷
        self.doc.add_heading('1. 초기 화면 (Screen Capture)', level=1)
        if os.path.exists(screenshot_path):
            try:
                # 페이지 폭에 맞춰 이미지 삽입 (약 6인치)
                self.doc.add_picture(screenshot_path, width=Inches(6.0))
            except Exception as e:
                self.doc.add_paragraph(f"[이미지 삽입 실패: {e}]")
        else:
            self.doc.add_paragraph("[스크린샷 파일을 찾을 수 없음]")

        # 3. 기능 및 구조 분석
        self.doc.add_heading('2. 기능 및 구조 분석', level=1)

        '''
        # 3-1. 주요 기능 추정
        self.doc.add_heading('2-1. 주요 기능 (추정)', level=2)
        features = page_data.get('guessed_features', [])
        if features:
            for feat in features:
                self.doc.add_paragraph(feat, style='List Bullet')
        else:
            self.doc.add_paragraph("특이 기능 발견되지 않음.")
            
        '''

        # 3-2. 구성 요소 요약
        self.doc.add_heading('구성 요소 요약', level=2)
        links = page_data.get('links', {})
        buttons = page_data.get('buttons', {})

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '항목'
        hdr_cells[1].text = '분석 결과'

        row_data = [
            ("Meta 설명", page_data.get('meta_description', '-')),
            ("총 링크 수", f"{links.get('total', 0)}개 (내부: {links.get('internal', 0)}, 외부: {links.get('external', 0)})"),
            ("버튼 수", f"{buttons.get('total', 0)}개 (Submit: {buttons.get('submit_type', 0)})"),
            ("입력 필드 수", f"{len(page_data.get('inputs', []))}개")
        ]

        for item, value in row_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(value)

        # 4. API 통신 분석
        self.doc.add_heading('3. API 통신 분석 (Network Logs)', level=1)
        self.doc.add_paragraph(f"수집된 주요 API 요청 수: {len(api_logs)}건")

        if api_logs:
            # 테이블 생성 (번호, Method, URL, IP, 상태)
            api_table = self.doc.add_table(rows=1, cols=4)
            api_table.style = 'Table Grid'
            api_table.autofit = False

            headers = api_table.rows[0].cells
            headers[0].text = 'Method'
            headers[1].text = 'URL (Path)'
            headers[2].text = 'Server IP'
            headers[3].text = 'Status'

            # 너무 많으면 상위 20개만 기록
            for log in api_logs[:20]:
                row_cells = api_table.add_row().cells
                row_cells[0].text = log['method']
                # URL이 너무 길면 잘라서 보여줌
                row_cells[1].text = (log['url'][:40] + '...') if len(log['url']) > 40 else log['url']
                row_cells[2].text = f"{log['domain']}\n({log['ip']})"
                row_cells[3].text = str(log['status'])
        else:
            self.doc.add_paragraph("감지된 XHR/Fetch 요청이 없습니다.")

        self.doc.add_page_break()
        self.doc.add_heading('4. 서버 IP WHOIS 정보', level=1)
        self.doc.add_paragraph("분석 과정에서 식별한 서버 IP의 소유자 정보 분석")

        if whois_data:
            #테이블 구조 : IP - 국가 - ISP - 할당 대역
            table = self.doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            hdr = table.rows[0].cells
            hdr[0].text = 'IP Address'
            hdr[1].text = '국가'
            hdr[2].text = '소유 기관'
            hdr[3].text = 'IP 대역'

            for cell in hdr :
                cell.width = Inches(1.5)

            for info in whois_data:
                row_cells = table.add_row().cells
                row_cells[0].text = info['ip']
                row_cells[1].text = info['country']
                row_cells[2].text = info['org']
                row_cells[3].text = info['cidr']

        else:
            self.doc.add_paragraph("조회된 WHOIS 정보 없음.")


        self.doc.add_page_break()  # 새 페이지에서 시작
        self.doc.add_heading('5. 상호작용 요소 상세 분석', level=1)
        self.doc.add_paragraph("사용자가 클릭 가능한 주요 요소들의 시각적 형태와 연결 동작을 분석합니다.")

        if interactive_elements:
            # 테이블 생성: [이미지] | [텍스트/유형] | [실행 동작]
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'

            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '요소 이미지'
            hdr_cells[1].text = '식별 정보'
            hdr_cells[2].text = '실행 동작 (Action)'

            # 열 너비 조정 (대략적으로)
            for cell in table.rows[0].cells:
                cell.width = Inches(2.0)

            for item in interactive_elements:
                row_cells = table.add_row().cells

                # 1열: 이미지 삽입
                p = row_cells[0].paragraphs[0]
                if os.path.exists(item['image_path']):
                    try:
                        run = p.add_run()
                        run.add_picture(item['image_path'], width=Inches(1.5))  # 너비 고정
                    except:
                        p.text = "[이미지 오류]"

                # 2열: 텍스트 및 태그 유형
                row_cells[1].text = f"Text: {item['text']}\nTag: <{item['type']}>"

                # 3열: 동작 정보 (URL 등)
                row_cells[2].text = item['action']

        else:
            self.doc.add_paragraph("분석 가능한 상호작용 요소를 찾지 못했습니다.")

        # 6. 전체 네트워크 리소스 분석
        if unique_ips or resource_summary:
            self.doc.add_page_break()
            self.doc.add_heading('6. 전체 네트워크 리소스 분석', level=1)
            self.doc.add_paragraph("웹사이트가 통신하는 모든 IP 주소와 리소스 타입을 분석합니다.")
            
            # 6-1. 리소스 타입별 통계
            if resource_summary:
                self.doc.add_heading('6-1. 리소스 타입별 통계', level=2)
                table = self.doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = '리소스 타입'
                hdr[1].text = '요청 수'
                
                for res_type, count in resource_summary.items():
                    row = table.add_row().cells
                    row[0].text = res_type
                    row[1].text = str(count)
            
            # 6-2. 고유 IP 주소 목록
            if unique_ips:
                self.doc.add_heading('6-2. 통신 대상 IP 주소 목록', level=2)
                self.doc.add_paragraph(f"총 {len(unique_ips)}개의 고유 IP 주소와 통신")
                
                table = self.doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = 'IP 주소'
                hdr[1].text = '도메인'
                hdr[2].text = '리소스 타입'
                hdr[3].text = '요청 수'
                
                import config
                max_ips = config.MAX_NETWORK_IPS
                
                for ip_info in unique_ips[:max_ips]:  # 설정된 최대 개수만 표시
                    row = table.add_row().cells
                    row[0].text = ip_info.get('ip', 'N/A')
                    row[1].text = ip_info.get('domain', 'N/A')
                    row[2].text = ', '.join(ip_info.get('types', []))
                    row[3].text = str(ip_info.get('count', 0))
                
                if len(unique_ips) > max_ips:
                    self.doc.add_paragraph(f"(총 {len(unique_ips)}개 중 상위 {max_ips}개만 표시)")

        # 7. OSINT 정보 (도메인 & SSL)
        if osint_data:
            self.doc.add_page_break()
            self.doc.add_heading('7. OSINT 정보 분석', level=1)
            
            # 7-1. 도메인 WHOIS 정보
            self.doc.add_heading('7-1. 도메인 WHOIS 정보', level=2)
            whois_info = osint_data.get('whois', {})
            
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            hdr = table.rows[0].cells
            hdr[0].text = '항목'
            hdr[1].text = '정보'
            
            whois_items = [
                ('도메인', osint_data.get('domain', 'N/A')),
                ('등록 기관', whois_info.get('registrar', 'N/A')),
                ('등록 일자', whois_info.get('creation_date', 'N/A')),
                ('만료 일자', whois_info.get('expiration_date', 'N/A')),
                ('조직', whois_info.get('org', 'N/A')),
            ]
            
            for item, value in whois_items:
                row = table.add_row().cells
                row[0].text = item
                row[1].text = str(value)
            
            # 네임서버
            name_servers = whois_info.get('name_servers', [])
            if name_servers:
                row = table.add_row().cells
                row[0].text = '네임 서버'
                row[1].text = '\n'.join(str(ns) for ns in name_servers)
            
            # 7-2. 서버 IP 주소 (nslookup)
            self.doc.add_heading('7-2. 서버 IP 주소 (nslookup)', level=2)
            server_ips = osint_data.get('server_ips', [])
            
            if server_ips:
                self.doc.add_paragraph(f"도메인 {osint_data.get('domain', 'N/A')}의 실제 서버 IP 주소:")
                
                table = self.doc.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = 'Server IP Address'
                
                for ip in server_ips:
                    row = table.add_row().cells
                    row[0].text = ip
            else:
                self.doc.add_paragraph("서버 IP 주소 조회 실패")
            
            # 7-3. DNS 레코드 (기타)
            self.doc.add_heading('7-3. DNS 레코드 (MX, TXT, NS)', level=2)
            dns_info = osint_data.get('dns', {})
            
            if any(dns_info.values()):
                for record_type, records in dns_info.items():
                    if records:
                        self.doc.add_paragraph(f"{record_type} 레코드:", style='List Bullet')
                        for record in records[:5]:  # 최대 5개
                            self.doc.add_paragraph(f"  • {record}", style='List Bullet 2')
            else:
                self.doc.add_paragraph("DNS 레코드 조회 실패")
            
            # 7-4. SSL/TLS 인증서
            self.doc.add_heading('7-4. SSL/TLS 인증서 정보', level=2)
            ssl_info = osint_data.get('ssl', {})
            
            if ssl_info.get('available'):
                table = self.doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = '항목'
                hdr[1].text = '정보'
                
                ssl_items = [
                    ('발급 기관', ssl_info.get('issuer', 'N/A')),
                    ('주체', ssl_info.get('subject', 'N/A')),
                    ('유효 시작', ssl_info.get('valid_from', 'N/A')),
                    ('유효 종료', ssl_info.get('valid_until', 'N/A')),
                ]
                
                for item, value in ssl_items:
                    row = table.add_row().cells
                    row[0].text = item
                    row[1].text = str(value)
                
                # SAN (Subject Alternative Names)
                san_list = ssl_info.get('san', [])
                if san_list:
                    row = table.add_row().cells
                    row[0].text = 'SAN (대체 이름)'
                    row[1].text = '\n'.join(san_list)
            else:
                self.doc.add_paragraph(f"SSL 인증서 없음: {ssl_info.get('reason', 'Unknown')}")

        # 8. 기술 스택 분석
        if osint_data and osint_data.get('technologies'):
            self.doc.add_page_break()
            self.doc.add_heading('8. 기술 스택 분석', level=1)
            
            tech = osint_data['technologies']
            
            # 프레임워크
            if tech.get('frameworks'):
                self.doc.add_heading('프론트엔드 프레임워크', level=2)
                for fw in tech['frameworks']:
                    self.doc.add_paragraph(f"• {fw}", style='List Bullet')
            
            # CMS
            if tech.get('cms'):
                self.doc.add_heading('CMS (콘텐츠 관리 시스템)', level=2)
                for cms in tech['cms']:
                    self.doc.add_paragraph(f"• {cms}", style='List Bullet')
            
            # 분석 도구
            if tech.get('analytics'):
                self.doc.add_heading('분석/추적 도구', level=2)
                for tool in tech['analytics']:
                    self.doc.add_paragraph(f"• {tool}", style='List Bullet')
            
            # 서버 정보
            if tech.get('server') and tech['server'] != 'Unknown':
                self.doc.add_heading('웹 서버', level=2)
                self.doc.add_paragraph(f"• {tech['server']}")
            
            # 프로그래밍 언어
            if tech.get('programming_language'):
                self.doc.add_heading('프로그래밍 언어', level=2)
                for lang in tech['programming_language']:
                    self.doc.add_paragraph(f"• {lang}", style='List Bullet')

        # 9. 연락처 정보
        if osint_data and osint_data.get('contacts'):
            import config
            
            contacts = osint_data['contacts']
            
            # 이메일이나 전화번호가 있을 때만 섹션 생성
            if contacts.get('emails') or contacts.get('phones'):
                self.doc.add_page_break()
                self.doc.add_heading('9. 연락처 정보', level=1)
                
                # 이메일
                if contacts.get('emails'):
                    self.doc.add_heading('발견된 이메일 주소', level=2)
                    for email in contacts['emails']:
                        self.doc.add_paragraph(f"• {email}", style='List Bullet')
                
                # 전화번호
                if contacts.get('phones'):
                    self.doc.add_heading('발견된 전화번호', level=2)
                    for phone in contacts['phones']:
                        self.doc.add_paragraph(f"• {phone}", style='List Bullet')
            
            # 소셜 미디어는 Section 10에만 표시 (중복 제거)
            # config.SHOW_SNS_IN_CONTACTS가 True일 때만 여기 표시
            if config.SHOW_SNS_IN_CONTACTS and contacts.get('social_media'):
                if not (contacts.get('emails') or contacts.get('phones')):
                    self.doc.add_page_break()
                    self.doc.add_heading('9. 연락처 정보', level=1)
                
                self.doc.add_heading('소셜 미디어 링크', level=2)
                for platform, links in contacts['social_media'].items():
                    self.doc.add_paragraph(f"{platform.capitalize()}:", style='List Bullet')
                    for link in links:
                        self.doc.add_paragraph(f"  • {link}", style='List Bullet 2')

        # 10. 심화 SNS 분석 (새로운 섹션)
        if sns_data:
            self.doc.add_page_break()
            self.doc.add_heading('10. 심화 소셜 미디어 분석', level=1)
            
            # 10-1. 요약 정보
            summary = sns_data.get('summary', {})
            self.doc.add_paragraph(f"총 {summary.get('total_links', 0)}개의 소셜 미디어 링크 발견")
            self.doc.add_paragraph(f"플랫폼 수: {summary.get('platform_count', 0)}개")
            self.doc.add_paragraph(f"소셜 미디어 존재감 점수: {sns_data.get('presence_score', 0)}/10")
            
            # 10-2. 플랫폼별 상세 정보
            links = sns_data.get('links', [])
            if links:
                self.doc.add_heading('10-1. 발견된 소셜 미디어 계정', level=2)
                
                # 테이블 생성
                table = self.doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = '플랫폼'
                hdr[1].text = '사용자명'
                hdr[2].text = '타입'
                hdr[3].text = '위치'
                hdr[4].text = '상태'
                hdr[5].text = 'URL'
                
                for link in links[:30]:  # 최대 30개
                    row = table.add_row().cells
                    row[0].text = link.get('platform', '').capitalize()
                    row[1].text = link.get('username', 'N/A')
                    row[2].text = link.get('type', 'N/A')
                    row[3].text = link.get('context', {}).get('location', 'N/A')
                    
                    # 검증 상태
                    validation = link.get('validation', {})
                    if validation:
                        if validation.get('is_valid'):
                            row[4].text = f"✓ {validation.get('status_code', '')}"
                        else:
                            row[4].text = "✗ 비활성"
                    else:
                        row[4].text = "미검증"
                    
                    # URL (짧게)
                    url_text = link.get('url', '')
                    row[5].text = url_text[:50] + '...' if len(url_text) > 50 else url_text
            
            # 10-3. 메타데이터 정보
            meta_tags = sns_data.get('meta_tags', {})
            if meta_tags:
                self.doc.add_heading('10-2. 소셜 미디어 메타데이터', level=2)
                
                table = self.doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                hdr = table.rows[0].cells
                hdr[0].text = '메타 태그'
                hdr[1].text = '값'
                
                for key, value in list(meta_tags.items())[:15]:  # 최대 15개
                    row = table.add_row().cells
                    row[0].text = key
                    row[1].text = str(value)[:100]
            
            # 10-4. 프로필 상세 정보 (메타데이터가 있는 경우)
            links_with_metadata = [l for l in links if l.get('metadata') and len(l.get('metadata', {})) > 1]
            if links_with_metadata:
                self.doc.add_heading('10-3. 프로필 상세 정보', level=2)
                
                for link in links_with_metadata[:10]:  # 최대 10개
                    platform = link.get('platform', '').capitalize()
                    username = link.get('username', 'N/A')
                    
                    self.doc.add_heading(f"{platform} - @{username}", level=3)
                    
                    metadata = link.get('metadata', {})
                    if 'error' not in metadata:
                        for key, value in metadata.items():
                            if value and str(value) != 'None':
                                self.doc.add_paragraph(f"• {key}: {value}", style='List Bullet')
                    else:
                        self.doc.add_paragraph(f"메타데이터 조회 실패: {metadata.get('error', 'Unknown')}")
            
            # 10-5. 분석 및 권장사항
            self.doc.add_heading('10-4. 분석 요약', level=2)
            
            presence_score = sns_data.get('presence_score', 0)
            platforms = summary.get('platforms', [])
            
            # 점수 기반 평가
            if presence_score >= 8:
                assessment = "매우 우수 - 강력한 소셜 미디어 존재감"
            elif presence_score >= 6:
                assessment = "우수 - 양호한 소셜 미디어 활용"
            elif presence_score >= 4:
                assessment = "보통 - 기본적인 소셜 미디어 존재"
            else:
                assessment = "미흡 - 소셜 미디어 활용 개선 필요"
            
            self.doc.add_paragraph(f"종합 평가: {assessment}")
            self.doc.add_paragraph(f"활용 플랫폼: {', '.join(platforms)}")
            
            # 링크 배치 분석
            footer_links = sum(1 for l in links if l.get('context', {}).get('location') == 'footer')
            header_links = sum(1 for l in links if l.get('context', {}).get('location') == 'header')
            
            self.doc.add_paragraph(f"• Footer 링크: {footer_links}개")
            self.doc.add_paragraph(f"• Header 링크: {header_links}개")
            
            # 검증 통계
            validated_links = [l for l in links if l.get('validation')]
            if validated_links:
                valid_count = sum(1 for l in validated_links if l.get('validation', {}).get('is_valid'))
                self.doc.add_paragraph(f"• 활성 링크: {valid_count}/{len(validated_links)}개")

    def save_file(self, filename):
        try:
            self.doc.save(filename)
            print(f"[Report] Document saved successfully: {filename}")
        except Exception as e:
            print(f"[Report] Error saving document: {e}")